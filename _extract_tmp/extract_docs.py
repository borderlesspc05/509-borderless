# -*- coding: utf-8 -*-
from pathlib import Path
import re
import sys

base = Path(r"C:\Users\Usuário\Downloads")
out_dir = Path(r"C:\Users\Usuário\OneDrive\BORDERLESS\509\_extract_tmp")
out_dir.mkdir(exist_ok=True)

def write_out(name, text):
    safe = re.sub(r'[<>:"/\\|?*]', "_", name)[:80]
    path = out_dir / f"{safe}.txt"
    path.write_text(text, encoding="utf-8")
    print(f"OK {path.name} ({len(text)} chars)")
    return path

def extract_docx(path):
    from docx import Document
    doc = Document(path)
    parts = [f"=== DOCX: {path.name} ===", f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}"]
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            style = p.style.name if p.style else ""
            parts.append(f"[P{i}|{style}] {t}")
    for ti, table in enumerate(doc.tables):
        parts.append(f"\n--- TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            # dedupe consecutive identical merged cells
            deduped = []
            for c in cells:
                if not deduped or deduped[-1] != c:
                    deduped.append(c)
            parts.append(f"R{ri}: " + " || ".join(deduped))
    return "\n".join(parts)

def extract_pdf_pdfplumber(path):
    import pdfplumber
    parts = [f"=== PDFplumber: {path.name} ==="]
    with pdfplumber.open(path) as pdf:
        parts.append(f"pages={len(pdf.pages)}")
        for i, page in enumerate(pdf.pages):
            parts.append(f"\n----- PAGE {i+1} -----")
            text = page.extract_text() or ""
            parts.append(text)
            tables = page.extract_tables() or []
            for ti, table in enumerate(tables):
                parts.append(f"\n--- TABLE p{i+1}-{ti} ---")
                for row in table:
                    cells = [(c or "").strip().replace("\n", " | ") for c in row]
                    parts.append(" || ".join(cells))
    return "\n".join(parts)

def extract_pdf_pymupdf(path):
    import fitz
    parts = [f"=== PyMuPDF: {path.name} ==="]
    doc = fitz.open(path)
    parts.append(f"pages={len(doc)}")
    for i, page in enumerate(doc):
        parts.append(f"\n----- PAGE {i+1} -----")
        parts.append(page.get_text("text"))
        # images count
        imgs = page.get_images(full=True)
        parts.append(f"[images_on_page={len(imgs)}]")
    return "\n".join(parts)

def count_replacement_chars(text):
    return text.count("\ufffd")

def extract_pdf_best(path):
    results = []
    try:
        t1 = extract_pdf_pdfplumber(path)
        results.append(("pdfplumber", t1, count_replacement_chars(t1)))
    except Exception as e:
        results.append(("pdfplumber", f"ERROR: {e}", 999999))
    try:
        t2 = extract_pdf_pymupdf(path)
        results.append(("pymupdf", t2, count_replacement_chars(t2)))
    except Exception as e:
        results.append(("pymupdf", f"ERROR: {e}", 999999))
    # pick fewer replacement chars, then longer
    results.sort(key=lambda x: (x[2], -len(x[1])))
    best = results[0]
    header = f"BEST={best[0]} replacement_chars={best[2]}\n"
    for name, text, rc in results:
        header += f"  alt {name}: chars={len(text)} replacement={rc}\n"
    return header + "\n" + best[1]

# Resolve POP filename (encoding)
pop = None
for p in base.glob("POP*.docx"):
    pop = p
    break

docx_files = [
    base / "Anamnese - musicoterapia (1).docx",
    pop,
    base / "Consideracoes setor TO 05-08-2026 - Novo Programa.docx",
]

pdf_files = [
    base / "Amiofe  (1).pdf",
    base / "ABFW_figuras_prova_fonologia (1).pdf",
    base / "GMFM-88-66_Translation-Portuguese-convertido (1) (1).pdf",
    base / "pediatric-balance-scale.pdf",
    base / "file.pdf",
]

print("=== DOCX ===")
for path in docx_files:
    if path is None or not path.exists():
        print("MISSING", path)
        continue
    try:
        text = extract_docx(path)
        write_out(path.stem, text)
    except Exception as e:
        print("FAIL DOCX", path, e)

print("=== PDF ===")
for path in pdf_files:
    if not path.exists():
        # try alternate Amiofe spacing
        alts = list(base.glob(path.name.split("(")[0].strip() + "*"))
        print("MISSING", path.name, "alts", [a.name for a in alts])
        if alts:
            path = alts[0]
        else:
            continue
    try:
        text = extract_pdf_best(path)
        write_out(path.stem, text)
    except Exception as e:
        print("FAIL PDF", path, e)

# Search Ashworth / agenda
print("\n=== SEARCH Ashworth / agenda ===")
patterns = re.compile(r"ashworth|status.?de.?agenda|escala.?modificada|ashwort", re.I)
hits = []
for p in base.rglob("*"):
    if not p.is_file():
        continue
    if patterns.search(p.name):
        hits.append(("filename", str(p)))
print("filename hits:", hits)

# Also scan text of extracted PDFs briefly for ashworth in file.pdf already
print("done")
